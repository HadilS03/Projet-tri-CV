import pandas as pd
from docx import Document # type: ignore

def generate_report(docx_file: str, user: str, results: list, weights: dict, retention_days: int, job_description_text: str):
    """Génère un rapport DOCX des résultats d'analyse des CV."""
    doc = Document()
    doc.add_heading('Rapport de tri de CV', 0)
    doc.add_paragraph(f"Recruteur : {user}")
    doc.add_paragraph(f"Date : {pd.Timestamp.now()}")
    doc.add_paragraph(f"Paramètres utilisés : Similarité={weights['sim']}, "
                      f"Mots-clés={weights['keywords']}, Expérience={weights['exp']}")
    doc.add_paragraph(f"Durée de rétention : {retention_days} jours")
    doc.add_heading("Description du poste", level=1)
    doc.add_paragraph(job_description_text)

    doc.add_heading('Résultats détaillés', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fichier'
    hdr_cells[1].text = 'Score (%)'
    hdr_cells[2].text = 'Similarité (%)'
    hdr_cells[3].text = 'Mots-clés (%)'
    hdr_cells[4].text = 'Expérience (ans)'

    for r in results:
        row_cells = table.add_row().cells
        row_cells[0].text = r.get('Fichier', '')
        row_cells[1].text = str(r.get("Score (%)", 0))
        row_cells[2].text = str(r.get("Similarité (%)", 0))
        row_cells[3].text = str(r.get("Mots-clés (%)", 0))
        row_cells[4].text = str(r.get("Expérience (ans)", 0))

    doc.save(docx_file)
