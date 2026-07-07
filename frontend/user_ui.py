import streamlit as st # type: ignore
import pandas as pd
from core.pipeline import process_cv
from core.auth import add_file
from docx import Document # type: ignore

# Utilitaires internes
def safe_float(val):
    if isinstance(val, dict):
        val = val.get("value", 0)
    try:
        return float(val)
    except:
        return 0

def generate_report(docx_file, user, results, weights, retention_days, job_description_text):
    doc = Document()
    doc.add_heading('Rapport de tri de CV', 0)
    doc.add_paragraph(f"Recruteur : {user}")
    doc.add_paragraph(f"Date : {pd.Timestamp.now()}")
    doc.add_paragraph(f"Paramètres utilisés : Similarité={weights['sim']}, "
                      f"Mots-clés={weights['keywords']}, Expérience={weights['exp']}")
    doc.add_paragraph(f"Durée de rétention : {retention_days} jours")
    doc.add_heading("Description du poste", level=1)
    doc.add_paragraph(job_description_text)

    doc.add_heading('Résultats détaillés', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fichier'
    hdr_cells[1].text = 'Score (%)'
    hdr_cells[2].text = 'Similarité (%)'
    hdr_cells[3].text = 'Mots-clés (%)'
    hdr_cells[4].text = 'Expérience (%)'

    for r in results:
        row_cells = table.add_row().cells
        row_cells[0].text = r['Fichier']
        row_cells[1].text = str(r['Score (%)'])
        row_cells[2].text = str(r['Similarité (%)'])
        row_cells[3].text = str(r['Mots-clés (%)'])
        row_cells[4].text = str(r['Expérience (%)'])

    doc.save(docx_file)


def user_app_ui():
    st.title("Tri automatique de CV")
    st.write(f"Bienvenue **{st.session_state.user}**")

    if st.button("Se déconnecter"):
        st.session_state.page = "login"
        st.session_state.user = None
        st.session_state.admin_logged = False

    # --- Paramètres de tri ---
    st.sidebar.subheader("⚙️ Paramètres de tri")
    w_sim = st.sidebar.slider("Poids Similarité", 0.0, 1.0, 0.7)
    w_keywords = st.sidebar.slider("Poids Mots-clés", 0.0, 1.0, 0.2)
    w_exp = st.sidebar.slider("Poids Expérience", 0.0, 1.0, 0.1)

    st.sidebar.subheader("Paramètres de gestion RGPD")
    retention_days = st.sidebar.selectbox(
        "Durée de rétention des données",
        options=[7, 15, 30],
        index=0,
        help="Les CV et les résultats obtenus seront automatiquement supprimés après ce délai."
    )

    # Normalisation des poids
    total_w = w_sim + w_keywords + w_exp
    if total_w > 0:
        w_sim /= total_w
        w_keywords /= total_w
        w_exp /= total_w

    # --- Description du poste ---
    st.subheader("📝 Description du poste")
    job_title = st.text_input("Intitulé du poste", placeholder="Ex: Data Scientist")
    job_skills = st.text_area("Compétences requises", placeholder="Python, SQL, Machine Learning…", height=80)
    job_tasks = st.text_area("Missions principales", placeholder="Analyser les données, construire des modèles…", height=80)
    job_experience = st.number_input("Expérience minimale (en années)", min_value=0, max_value=50, value=2)

    job_description_text = f"{job_title}\nCompétences: {job_skills}\nMissions: {job_tasks}\nExpérience: {job_experience} ans"

    consent = st.checkbox("Confirmer le consentement au RGPD.")

    # --- Upload de CV ---
    st.subheader("📂 Importez les CV")
    uploaded_files = st.file_uploader(
        "Sélectionnez un ou plusieurs fichiers",
        type=["pdf", "docx"],
        accept_multiple_files=True
    )

    # --- Analyse ---
    if st.button("Lancer l'analyse"):
        if not job_title or not job_skills or not job_tasks or not uploaded_files or not consent:
            st.error("Remplissez toutes les informations et cochez le consentement.")
            st.stop()

        results = []
        with st.spinner("Analyse en cours…"):
            for uf in uploaded_files:
                temp_path = f"temp_{uf.name}"
                with open(temp_path, "wb") as f:
                    f.write(uf.getbuffer())
                try:
                    res = process_cv(temp_path, job_description_text, w_sim, w_keywords, w_exp)
                    results.append({
                        "Fichier": uf.name,
                        "Score (%)": safe_float(res.get("score", 0)),
                        "Similarité (%)": safe_float(res.get("similarity", 0)),
                        "Mots-clés (%)": safe_float(res.get("keywords", 0)),
                        "Expérience (%)": safe_float(res.get("experience", 0))
                    })
                    add_file(uf.name, retention_days)
                except Exception as e:
                    st.error(f"Erreur avec {uf.name} : {e}")

        results = sorted(results, key=lambda x: x["Score (%)"], reverse=True)

        if results:
            df = pd.DataFrame(results)
            st.subheader("📊 Résultats")
            st.dataframe(df, use_container_width=True)

            # Export CSV
            csv = df.to_csv(index=False).encode("utf-8")
            st.download_button("⬇️ Télécharger CSV", data=csv, file_name="resultats_cv.csv", mime="text/csv")

            # Rapport DOCX
            report_file = "rapport_cv.docx"
            weights = {"sim": w_sim, "keywords": w_keywords, "exp": w_exp}
            generate_report(report_file, st.session_state.user, results, weights, retention_days, job_description_text)

            st.download_button(
                "⬇️ Télécharger le rapport DOCX",
                data=open(report_file, "rb").read(),
                file_name="rapport_cv.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
