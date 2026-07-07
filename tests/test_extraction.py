"""Tests de l'extraction de texte et du rejet des formats non supportés."""
import pytest
from docx import Document  # type: ignore

import core.pipeline as p


def test_extract_text_from_docx(tmp_path):
    chemin = tmp_path / "cv.docx"
    doc = Document()
    doc.add_paragraph("Developpeur Python")
    doc.add_paragraph("5 ans d'experience")
    doc.save(str(chemin))

    texte = p.extract_text_from_docx(str(chemin))
    assert "Developpeur Python" in texte
    assert "experience" in texte


def test_extension_inconnue_leve_valueerror():
    # process_cv doit rejeter une extension non gérée avant tout traitement.
    with pytest.raises(ValueError):
        p.process_cv("fichier.txt", "annonce")
